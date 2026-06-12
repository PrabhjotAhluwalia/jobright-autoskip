from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX_PATH = DOCS / "Navjeet_JobRight_Mac_Guide.docx"

NAVY = "17324D"
BLUE = "2867B2"
TEAL = "168C83"
PALE_BLUE = "EAF3FB"
PALE_GREEN = "EAF7F3"
PALE_GOLD = "FFF4D6"
PALE_RED = "FDEBEC"
RED = "A3262A"
GRAY = "5B6573"
LIGHT_GRAY = "D9DEE5"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=95, start=150, bottom=95, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LIGHT_GRAY, size=8):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=None, bold=None, color=None, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_font(r, 22, True, NAVY)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(10)
        r2 = p2.add_run(subtitle)
        set_font(r2, 11.5, False, GRAY)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text, checked=False):
    p = doc.add_paragraph(style="List Bullet")
    if checked:
        lead = p.add_run("CHECK: ")
        set_font(lead, bold=True, color=TEAL)
    run = p.add_run(text)
    set_font(run)
    return p


def add_step(doc, number, title, detail):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [650, 8710])
    set_table_borders(table, color="FFFFFF", size=0)
    ncell, dcell = table.rows[0].cells
    set_cell_shading(ncell, TEAL)
    set_cell_shading(dcell, PALE_GREEN)
    for cell in (ncell, dcell):
        set_cell_margins(cell, 75, 140, 75, 140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = ncell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(number))
    set_font(r, 12.5, True, WHITE)
    p2 = dcell.paragraphs[0]
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(title)
    set_font(r2, 10, True, NAVY)
    p3 = dcell.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run(detail)
    set_font(r3, 9, False, GRAY)


def add_callout(doc, label, text, fill=PALE_BLUE, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size=4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 180, 120, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label.upper())
    set_font(r1, 9, True, label_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_font(r2, 9.5, False, NAVY)


def add_command(doc, command):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="CDD5DF", size=6)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7F9")
    set_cell_margins(cell, 75, 150, 75, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(command)
    set_font(r, 8.5, False, NAVY, name="Menlo")


def add_checklist_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [650, 2850, 5860])
    set_table_borders(table)
    headers = ["Done", "Item", "What to verify"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_font(r, 10, True, WHITE)
    for item, detail in rows:
        cells = table.add_row().cells
        cells[0].text = "[  ]"
        cells[1].text = item
        cells[2].text = detail
        for i, cell in enumerate(cells):
            set_cell_margins(cell, 75, 120, 75, 120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                set_font(run, 8.5, bold=(i == 1), color=NAVY if i == 1 else GRAY)
    return table


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("JobRight Auto-Skip | Private setup guide for Navjeet")
    set_font(r, 7.5, False, GRAY)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Heading 1", 15, NAVY, 9, 5),
        ("Heading 2", 11.5, BLUE, 7, 4),
        ("Heading 3", 10.5, TEAL, 6, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(9.5)
        style.paragraph_format.left_indent = Inches(0.34)
        style.paragraph_format.first_line_indent = Inches(-0.17)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.08


def page_break(doc):
    doc.add_page_break()


def build():
    DOCS.mkdir(exist_ok=True)
    doc = Document()
    configure_styles(doc)
    props = doc.core_properties
    props.title = "JobRight Auto-Skip: Mac Setup Guide"
    props.subject = "Private onboarding and profile-personalization guide"
    props.author = "JobRight Auto-Skip repository owner"

    add_title(
        doc,
        "JobRight Auto-Skip",
        "Simple Mac setup and safe-use guide for Navjeet | Software Engineer edition",
    )
    add_callout(
        doc,
        "Read this first",
        "This copy initially contains another person's resume, identity, job preferences, and application answers. Do not start a job queue until every red item in this guide has been replaced with your own truthful information.",
        fill=PALE_RED,
        label_color=RED,
    )
    add_heading(doc, "The whole process", 1)
    for number, title, detail in (
        (1, "Accept the private GitHub invitation", "Use the invited Georgia Tech GitHub account."),
        (2, "Download the extension", "Clone the private repository to your Mac."),
        (3, "Replace the owner's personal data", "Use your resume, contact details, work authorization, sponsorship, location, and salary preferences."),
        (4, "Describe your software-engineering profile", "Add only skills, years, projects, and achievements that are true."),
        (5, "Load it in Chrome", "Use Chrome's Extensions page and select the downloaded folder."),
        (6, "Run one supervised test", "Pause automatic submission and verify every answer before hands-free use."),
    ):
        add_step(doc, number, title, detail)

    add_heading(doc, "Before you begin", 1)
    add_bullet(doc, "A MacBook with Google Chrome installed")
    add_bullet(doc, "The invited Georgia Tech GitHub account")
    add_bullet(doc, "Your current software-engineering resume as a PDF")
    add_bullet(doc, "Your LinkedIn URL, phone number, location, and salary target")
    add_bullet(doc, "Your exact work-authorization and sponsorship answers")
    add_callout(
        doc,
        "Private access",
        "Only the repository owner and invited GitHub collaborators can see this repository. Do not forward the files or invite anyone else.",
        fill=PALE_GOLD,
        label_color="8A6500",
    )

    page_break(doc)
    add_title(doc, "1. Get the private files", "Accept the invitation, then download the repository.")
    add_heading(doc, "Accept the invitation", 1)
    add_step(doc, 1, "Open the GitHub email", "Look for an invitation to PrabhjotAhluwalia/jobright-autoskip.")
    add_step(doc, 2, "Sign in to GitHub", "Use the invited Georgia Tech GitHub account.")
    add_step(doc, 3, "Click Accept invitation", "The repository will remain private after you accept.")
    add_heading(doc, "Download it with Terminal", 1)
    add_body(doc, "Open Terminal: press Command + Space, type Terminal, then press Return.")
    add_body(doc, "Paste these commands one at a time and press Return after each:")
    add_command(doc, "xcode-select --install")
    add_command(doc, "cd ~/Documents")
    add_command(doc, "git clone https://github.com/PrabhjotAhluwalia/jobright-autoskip.git")
    add_command(doc, "cd jobright-autoskip")
    add_command(doc, "git switch -c navjeet/profile")
    add_callout(
        doc,
        "If GitHub asks for a password",
        "Sign in through the browser prompt or use a GitHub personal access token. Your normal GitHub password may not work for command-line cloning.",
    )
    add_heading(doc, "Where the folder should be", 1)
    add_body(doc, "Finder > Documents > jobright-autoskip")

    page_break(doc)
    add_title(doc, "2. Replace all personal information", "These changes are mandatory before the first application.")
    add_heading(doc, "A. Replace the profile prompt", 1)
    add_body(doc, "Open jobright_system_prompt.txt in a code editor. Replace every fact about the current owner.")
    add_checklist_table(
        doc,
        [
            ("Identity", "Full name, first name, last name, email, phone, LinkedIn, GitHub"),
            ("Location", "Current city, commute distance, relocation, remote, hybrid, and onsite preferences"),
            ("Legal answers", "Country, work authorization, sponsorship, clearance, and availability"),
            ("Career", "Job titles, employers, dates, years of experience, projects, and achievements"),
            ("Education", "School, degree, graduation date, certifications, and relevant coursework"),
            ("Compensation", "Minimum salary, target range, and whether the answer is base or total compensation"),
        ],
    )
    add_heading(doc, "B. Replace the contact constants", 1)
    add_body(doc, "Open ats_content.js and find PROFILE_CORRECTIONS. Replace every value:")
    add_command(doc, "fullName, firstName, lastName, email, phone, linkedin")
    add_callout(
        doc,
        "Phone format",
        "Use digits only in the phone constant. Example: 4045551234. Do not include spaces, parentheses, or +1.",
    )
    add_heading(doc, "C. Replace the resume", 1)
    add_step(doc, 1, "Remove the existing resume", "Delete the other candidate's PDF from the assets folder.")
    add_step(doc, 2, "Add your resume", "Use a clear filename such as Navjeet_Software_Engineer_Resume.pdf.")
    add_step(doc, 3, "Update ats_content.js", "Change FALLBACK_RESUME_PATH and FALLBACK_RESUME_NAME to your filename.")
    add_step(doc, 4, "Update manifest.json", "Replace the old resume filename under web_accessible_resources.")

    page_break(doc)
    add_title(doc, "3. Build your software-engineer profile", "Use facts from your resume and real experience, not generic claims.")
    add_heading(doc, "Choose the roles you actually want", 1)
    add_bullet(doc, "Software Engineer / Software Developer")
    add_bullet(doc, "Backend Engineer")
    add_bullet(doc, "Frontend Engineer")
    add_bullet(doc, "Full Stack Engineer")
    add_bullet(doc, "Platform / Infrastructure / DevOps Engineer")
    add_bullet(doc, "Data Engineer or ML Engineer, only if your background supports it")
    add_heading(doc, "Fill in these technical fields", 1)
    add_checklist_table(
        doc,
        [
            ("Languages", "Java, Python, JavaScript/TypeScript, C++, C#, Go, Rust, or others; include truthful years"),
            ("Frontend", "React, Next.js, Angular, Vue, HTML/CSS, browser APIs, accessibility"),
            ("Backend", "Node.js, Spring Boot, Django, FastAPI, REST, GraphQL, gRPC, microservices"),
            ("Databases", "PostgreSQL, MySQL, SQL Server, MongoDB, Redis, DynamoDB, warehouses"),
            ("Cloud", "AWS, Azure, GCP, serverless, Docker, Kubernetes"),
            ("DevOps", "GitHub Actions, CI/CD, Terraform, monitoring, logging, incident response"),
            ("Testing", "Unit, integration, end-to-end, performance, security, test frameworks"),
            ("System design", "Distributed systems, APIs, caching, queues, scalability, reliability"),
            ("AI/ML", "Models, frameworks, data pipelines, evaluation, or LLM work only if real"),
            ("Projects", "Problem, your contribution, technologies, scale, and measurable result"),
        ],
    )
    add_heading(doc, "Use specific evidence", 1)
    add_body(doc, "Good: “Built a Spring Boot API used by 20,000 monthly users and reduced p95 latency by 35%.”")
    add_body(doc, "Avoid: “Expert in every programming language and cloud platform.”")
    add_callout(
        doc,
        "Years of experience",
        "Do not give every skill the same number of years. Use actual professional, internship, academic, or project experience and label it accurately.",
        fill=PALE_GOLD,
        label_color="8A6500",
    )

    add_title(doc, "4. Review the automatic answers", "The extension must match your real situation.")
    add_heading(doc, "Current automatic rules", 1)
    add_checklist_table(
        doc,
        [
            ("Sponsorship", "Currently answers No. Change it if you need sponsorship now or in the future."),
            ("Work authorization", "Currently answers Yes. Change it if you are not authorized for that country or employer."),
            ("Onsite / hybrid", "Currently answers Yes to onsite, hybrid, commute, local, and relocation questions."),
            ("Country", "Currently defaults empty country fields to United States."),
            ("Recruiter SMS", "Currently answers No to recruiter text-message opt-in."),
            ("Consent", "May accept required terms, privacy, and consent checkboxes."),
            ("Resume", "May upload the fallback PDF when a site needs a resume."),
        ],
    )
    add_callout(
        doc,
        "Never guess legal answers",
        "Authorization, sponsorship, citizenship, clearance, criminal-history, disability, veteran, demographic, and conflict-of-interest questions must be answered truthfully. Pause instead of guessing.",
        fill=PALE_RED,
        label_color=RED,
    )
    add_heading(doc, "Location questions", 1)
    add_body(doc, "The extension may answer Yes to questions such as “Can you work in our office five days a week?” Only keep that behavior if you can actually commute or relocate.")
    add_heading(doc, "Salary questions", 1)
    add_body(doc, "Choose a realistic minimum and target range. State whether it is annual base salary or total compensation. Review unusual hourly, contract, or international-currency questions manually.")
    add_heading(doc, "Job filters", 1)
    add_body(doc, "Review these files so the queue targets software-engineering roles:")
    add_command(doc, "jobright_excluded_job_titles.txt")
    add_command(doc, "jobright_excluded_job_title_regexes.txt")
    add_command(doc, "shared_blocklist.json")
    add_body(doc, "Consider excluding sales, marketing, product management, support, hardware, QA-only, director/VP, and mismatched seniority roles.")

    page_break(doc)
    add_title(doc, "5. Install it in Chrome", "Load the downloaded folder as a private unpacked extension.")
    for number, title, detail in (
        (1, "Open Chrome", "Use Google Chrome, not Safari."),
        (2, "Open the Extensions page", "Type chrome://extensions in the address bar and press Return."),
        (3, "Turn on Developer mode", "Use the switch in the upper-right corner."),
        (4, "Click Load unpacked", "A folder picker will open."),
        (5, "Select jobright-autoskip", "Choose Finder > Documents > jobright-autoskip, then click Select."),
        (6, "Pin the extension", "Open Chrome's puzzle-piece menu and pin JobRight Auto-Skip."),
        (7, "Open JobRight", "Sign in and confirm the extension popup says it is active."),
    ):
        add_step(doc, number, title, detail)
    add_heading(doc, "After changing any profile file", 1)
    add_step(doc, 1, "Return to chrome://extensions", "Find JobRight Auto-Skip.")
    add_step(doc, 2, "Click Reload", "Use the circular-arrow button on the extension card.")
    add_step(doc, 3, "Refresh JobRight", "Refresh JobRight and every open application tab.")
    add_callout(
        doc,
        "Gmail codes are optional",
        "The owner's Google OAuth secret is not included. Start without automatic Gmail OTP. Configure your own Google Cloud OAuth application later if needed. Never commit passwords, tokens, client secrets, or one-time codes.",
    )

    page_break(doc)
    add_title(doc, "6. Run one safe test", "Watch the first application from beginning to end.")
    add_callout(
        doc,
        "Keep auto-submit paused",
        "Use one test application. Do not turn on hands-free operation until every field below is correct and the application has not been submitted.",
        fill=PALE_GOLD,
        label_color="8A6500",
    )
    add_checklist_table(
        doc,
        [
            ("Name", "Your name appears exactly as intended"),
            ("Contact", "Your email, phone, LinkedIn, and GitHub are correct"),
            ("Resume", "Your software-engineering resume is attached"),
            ("Location", "City, country, commute, onsite, hybrid, and relocation answers are true"),
            ("Legal", "Work authorization and sponsorship are correct"),
            ("Compensation", "Salary amount and currency are appropriate"),
            ("Experience", "Years, employers, education, skills, and projects are factual"),
            ("Options", "Selected radio buttons and checkboxes stay selected"),
            ("Submission", "A successful application is counted once; failures are not counted"),
            ("Timeout", "A genuinely stuck job saves a screenshot and skips after 100 seconds without progress"),
            ("Prompt cycle", "The system prompt appears only after 10 confirmed successful submissions"),
        ],
    )
    add_heading(doc, "When to stop immediately", 1)
    add_bullet(doc, "A legal, sponsorship, authorization, or location answer is wrong")
    add_bullet(doc, "The extension selects an option and then removes it")
    add_bullet(doc, "The wrong resume or contact information appears")
    add_bullet(doc, "A site reports success but JobRight remains stuck")
    add_bullet(doc, "The same application is counted more than once")
    add_body(doc, "Pause the queue, take a full screenshot showing the question and all options, and fix the rule before continuing.")

    page_break(doc)
    add_title(doc, "7. Quick help", "The most common fixes.")
    add_heading(doc, "Changes are not appearing", 1)
    add_body(doc, "Reload the extension at chrome://extensions, then refresh JobRight and every application tab.")
    add_heading(doc, "Resume upload fails", 1)
    add_body(doc, "Confirm the PDF filename matches in all three places: the assets folder, ats_content.js, and manifest.json.")
    add_heading(doc, "A wrong Yes or No is selected", 1)
    add_body(doc, "Pause immediately. Capture the full question and every answer choice. Do not continue until the matching rule is corrected.")
    add_heading(doc, "GitHub updates are available", 1)
    add_command(doc, "cd ~/Documents/jobright-autoskip")
    add_command(doc, "git switch navjeet/profile")
    add_command(doc, "git fetch origin")
    add_command(doc, "git rebase origin/main")
    add_callout(
        doc,
        "Protect your profile",
        "During updates, do not overwrite your personalized prompt, resume, PROFILE_CORRECTIONS, authorization answers, sponsorship answers, or location preferences with the owner's values.",
        fill=PALE_RED,
        label_color=RED,
    )
    add_heading(doc, "Final ready-to-run checklist", 1)
    add_checklist_table(
        doc,
        [
            ("Private invitation", "Accepted using the correct GitHub account"),
            ("Personal prompt", "Every owner-specific fact replaced"),
            ("Contact details", "PROFILE_CORRECTIONS uses your data"),
            ("Resume", "Your PDF filename matches in all required files"),
            ("Legal answers", "Authorization and sponsorship are truthful"),
            ("Preferences", "Location, onsite, hybrid, relocation, salary, and role filters are correct"),
            ("Chrome reload", "Extension and application pages refreshed"),
            ("Supervised test", "One complete application reviewed before hands-free use"),
        ],
    )
    add_body(doc, "Repository: https://github.com/PrabhjotAhluwalia/jobright-autoskip")

    for section in doc.sections:
        add_footer(section)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
